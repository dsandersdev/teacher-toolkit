from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocxExporter:
    def __init__(self, output_dir: str = "outputs"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)

    def save(
        self,
        content: str,
        prefix: str,
        metadata: dict | None = None,
    ) -> Path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = self.output_dir / f"{prefix}_{timestamp}.docx"

        metadata = metadata or {}

        document = Document()

        # Page setup
        section = document.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        # Fonts
        styles = document.styles
        styles["Normal"].font.name = "Calibri"
        styles["Normal"].font.size = Pt(11)

        title = metadata.get(
            "title",
            prefix.replace("_", " ").title(),
        )

        # Title
        heading = document.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata section
        if metadata:
            meta_table = document.add_table(rows=0, cols=2)

            for key, value in metadata.items():
                row = meta_table.add_row()
                row.cells[0].text = key.replace("_", " ").title()
                row.cells[1].text = str(value)

            document.add_paragraph("")

        # Content
        for line in content.splitlines():
            clean_line = line.strip()

            if not clean_line:
                continue

            if clean_line.startswith("# "):
                document.add_heading(
                    clean_line[2:],
                    level=1,
                )

            elif clean_line.startswith("## "):
                document.add_heading(
                    clean_line[3:],
                    level=2,
                )

            elif clean_line.startswith("### "):
                document.add_heading(
                    clean_line[4:],
                    level=3,
                )

            elif clean_line.startswith("- "):
                document.add_paragraph(
                    clean_line[2:],
                    style="List Bullet",
                )

            elif (
                len(clean_line) > 2
                and clean_line[0].isdigit()
                and clean_line[1:3] in [". ", ") "]
            ):
                document.add_paragraph(
                    clean_line,
                    style="List Number",
                )

            else:
                paragraph = document.add_paragraph(clean_line)
                paragraph.paragraph_format.space_after = Pt(8)

        # Footer
        footer = document.sections[0].footer
        footer.paragraphs[0].text = (
            "Generated with Teacher Toolkit"
        )
        footer.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        document.save(output_path)

        return output_path